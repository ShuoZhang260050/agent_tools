import io
import os
from unittest.mock import patch
from agent.tools.read_file import ReadFileTool


def _run_with_ws(file_path, ws_path, **kwargs):
    user_id = kwargs.pop("user_id", 1)
    with patch("agent.memory.user_memory.get_workspace", return_value=ws_path):
        tool = ReadFileTool()
        return tool._run(file_path, config={"configurable": {"user_id": user_id}}, **kwargs)


def test_read_text_file(tmp_path):
    f = tmp_path / "hello.py"
    f.write_text("print('hello')\nprint('world')\n", encoding="utf-8")
    out = _run_with_ws("hello.py", str(tmp_path))
    assert "hello" in out
    assert "world" in out
    assert "<external_content" in out
    assert "2 行" in out


def test_read_line_range(tmp_path):
    lines = [f"line {i}" for i in range(1, 11)]
    f = tmp_path / "data.txt"
    f.write_text("\n".join(lines) + "\n", encoding="utf-8")
    out = _run_with_ws("data.txt", str(tmp_path), start_line=3, end_line=5)
    assert "line 3" in out
    assert "line 5" in out
    assert "line 6" not in out
    assert "start_line=6" in out


def test_read_bad_ext(tmp_path):
    f = tmp_path / "blob.bin"
    f.write_bytes(b"\x00\x01\x02")
    out = _run_with_ws("blob.bin", str(tmp_path))
    assert "不支持" in out


def test_read_blocked_env(tmp_path):
    f = tmp_path / ".env"
    f.write_text("SECRET=xxx", encoding="utf-8")
    out = _run_with_ws(".env", str(tmp_path))
    assert "拒绝" in out or "敏感" in out


def test_read_path_traversal(tmp_path):
    sub = tmp_path / "sub"
    sub.mkdir()
    out = _run_with_ws("../../etc/passwd", str(tmp_path))
    assert "超出工作空间" in out or "不存在" in out


def test_read_no_workspace(tmp_path):
    with patch("agent.memory.user_memory.get_workspace", return_value=None):
        tool = ReadFileTool()
        out = tool._run("test.py", config={"configurable": {"user_id": 1}})
        assert "未设置工作空间" in out


def test_read_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello from docx")
    doc.add_paragraph("Second line")
    f = tmp_path / "test.docx"
    doc.save(str(f))
    out = _run_with_ws("test.docx", str(tmp_path))
    assert "Hello from docx" in out
    assert "Second line" in out


def test_read_file_accepts_path_keyword(tmp_path):
    """回归：LLM 常用 path= 调用 read_file（与 write_file/edit_file 一致），
    read_file 的参数必须叫 path 而非 file_path。"""
    f = tmp_path / "kw.py"
    f.write_text("print('ok')\n", encoding="utf-8")
    with patch("agent.memory.user_memory.get_workspace", return_value=str(tmp_path)):
        tool = ReadFileTool()
        out = tool._run(path="kw.py", config={"configurable": {"user_id": 1}})
    assert "ok" in out
